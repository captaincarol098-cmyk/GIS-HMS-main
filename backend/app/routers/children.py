from datetime import date
from uuid import UUID
from fastapi import APIRouter, Depends, HTTPException, Query, Request
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload
from ..database import get_db
from ..middleware.rbac import assert_barangay_scope, get_current_user, scoped_barangay_filter
from ..models import Alert, Child, Measurement, User
from ..models.entities import Sex
from ..schemas.common import ChildCreate
from ..services.audit import log_activity
from ..services.websocket import manager

router = APIRouter(prefix="/api/children", tags=["children"])


def age_months(birth_date: date, on_date: date | None = None) -> int:
    on_date = on_date or date.today()
    return max(0, (on_date.year - birth_date.year) * 12 + on_date.month - birth_date.month - (1 if on_date.day < birth_date.day else 0))


def child_payload(child: Child) -> dict:
    latest = sorted(child.measurements, key=lambda m: m.measurement_date, reverse=True)[0] if child.measurements else None
    return {
        "id": str(child.id),
        "full_name": child.full_name,
        "birth_date": child.birth_date,
        "age_months": age_months(child.birth_date),
        "sex": child.sex.value,
        "guardian_name": child.guardian_name,
        "contact_number": child.contact_number,
        "purok_id": str(child.purok_id),
        "barangay_id": str(child.barangay_id),
        "latitude": child.latitude,
        "longitude": child.longitude,
        "is_active": child.is_active,
        "latest_measurement": None if not latest else {"id": str(latest.id), "date": latest.measurement_date, "overall_status": latest.overall_status.value, "waz": latest.waz, "haz": latest.haz, "whz": latest.whz},
    }


@router.get("")
async def list_children(search: str | None = None, barangay_id: UUID | None = None, purok_id: UUID | None = None, sex: str | None = None, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    assert_barangay_scope(user, barangay_id)
    stmt = select(Child).options(selectinload(Child.measurements)).where(Child.is_active.is_(True)).order_by(Child.full_name)
    stmt = scoped_barangay_filter(stmt, Child, user)
    if barangay_id:
        stmt = stmt.where(Child.barangay_id == barangay_id)
    if purok_id:
        stmt = stmt.where(Child.purok_id == purok_id)
    if sex:
        stmt = stmt.where(Child.sex == Sex(sex))
    if search:
        stmt = stmt.where(or_(Child.full_name.ilike(f"%{search}%"), Child.guardian_name.ilike(f"%{search}%")))
    return [child_payload(c) for c in (await db.scalars(stmt)).all()]


@router.post("")
async def create_child(body: ChildCreate, request: Request, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    assert_barangay_scope(user, body.barangay_id)
    child = Child(**body.model_dump())
    db.add(child)
    await db.flush()
    await log_activity(db, user.id, "CREATE_CHILD", "children", str(child.id), body.model_dump(mode="json"), request.client.host if request.client else None)
    await db.commit()
    await manager.broadcast("refetch_data")
    await db.refresh(child)
    return child_payload(child)


@router.get("/{child_id}")
async def get_child(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    child = await db.scalar(select(Child).options(selectinload(Child.measurements)).where(Child.id == child_id))
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    alerts = (await db.scalars(select(Alert).where(Alert.child_id == child_id, Alert.is_resolved.is_(False)))).all()
    payload = child_payload(child)
    payload["measurements"] = [m for m in sorted(child.measurements, key=lambda x: x.measurement_date, reverse=True)]
    payload["active_alerts"] = [{"id": str(a.id), "severity": a.severity.value, "message": a.message} for a in alerts]
    return payload


@router.put("/{child_id}")
async def update_child(child_id: UUID, body: ChildCreate, request: Request, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    for key, value in body.model_dump().items():
        setattr(child, key, value)
    await log_activity(db, user.id, "UPDATE_CHILD", "children", str(child.id), {}, request.client.host if request.client else None)
    await db.commit()
    await manager.broadcast("refetch_data")
    await db.refresh(child)
    return child_payload(child)


@router.delete("/{child_id}")
async def delete_child(child_id: UUID, request: Request, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    child.is_active = False
    await log_activity(db, user.id, "DELETE_CHILD", "children", str(child.id), {}, request.client.host if request.client else None)
    await db.commit()
    await manager.broadcast("refetch_data")
    return {"ok": True}


@router.get("/{child_id}/measurements")
async def child_measurements(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    return list((await db.scalars(select(Measurement).where(Measurement.child_id == child_id).order_by(Measurement.measurement_date.desc()))).all())


@router.get("/{child_id}/growth-chart-data")
async def growth_chart_data(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    rows = await child_measurements(child_id, db, user)
    return {"measurements": rows, "reference_lines": [{"age_months": i, "sd3neg": -3, "sd2neg": -2, "sd0": 0, "sd2pos": 2, "sd3pos": 3} for i in range(0, 61)]}


@router.get("/{child_id}/rag-recommendations")
async def rag_recommendations(
    child_id: UUID,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user)
):
    """
    Generate evidence-based recommendations using RAG (Retrieval-Augmented Generation).
    
    RAG Steps:
    1. Context Retrieval - Get child's data (age, sex, measurements, growth trends)
    2. Knowledge Retrieval - Query knowledge base (WHO guidelines, NNC protocols, best practices)
    3. Recommendation Generation - Create tailored, evidence-based recommendations
    4. Risk Prioritization - Calculate risk score based on multiple factors
    5. Output Delivery - Format for dashboard display
    
    Returns comprehensive recommendations with evidence citations.
    """
    from ..services.rag_recommendations import generate_child_recommendations
    
    # Check child exists and user has access
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    
    # BHW can only access children in their barangay
    if user.role.value == "admin" and child.barangay_id != user.barangay_id:
        raise HTTPException(403, "Access denied - child not in your barangay")
    
    recommendations = await generate_child_recommendations(child_id, db)
    
    return recommendations


@router.get("/{child_id}/download/growth-data")
async def download_growth_data(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    """Download growth data as CSV"""
    from fastapi.responses import StreamingResponse
    import csv
    import io
    
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    
    measurements = list((await db.scalars(
        select(Measurement).where(Measurement.child_id == child_id).order_by(Measurement.measurement_date)
    )).all())
    
    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Weight (kg)", "Height (cm)", "MUAC (cm)", "WAZ", "HAZ", "WHZ", "Status"])
    
    for m in measurements:
        writer.writerow([
            m.measurement_date.isoformat(),
            m.weight_kg,
            m.height_cm,
            m.muac_cm or "",
            m.waz,
            m.haz,
            m.whz,
            m.overall_status.value if m.overall_status else ""
        ])
    
    # Convert to bytes
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={child.full_name}_growth_data.csv"}
    )


@router.get("/{child_id}/download/health-entries")
async def download_health_entries(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    """Download health entries/observations as CSV"""
    from fastapi.responses import StreamingResponse
    import csv
    import io
    
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    
    measurements = list((await db.scalars(
        select(Measurement).where(Measurement.child_id == child_id).order_by(Measurement.measurement_date)
    )).all())
    
    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Overall Status", "Observations", "Nutrition Level", "Actions Taken"])
    
    for m in measurements:
        writer.writerow([
            m.measurement_date.isoformat(),
            m.overall_status.value if m.overall_status else "",
            m.notes or "",
            m.nutritional_status if hasattr(m, 'nutritional_status') else "",
            m.intervention if hasattr(m, 'intervention') else ""
        ])
    
    # Convert to bytes
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={child.full_name}_health_entries.csv"}
    )


@router.get("/{child_id}/download/complete-report")
async def download_complete_report(child_id: UUID, db: AsyncSession = Depends(get_db), user: User = Depends(get_current_user)):
    """Download complete child report as Word document"""
    from fastapi.responses import StreamingResponse
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import io
    
    child = await db.get(Child, child_id)
    if not child:
        raise HTTPException(404, "Child not found")
    assert_barangay_scope(user, child.barangay_id)
    
    measurements = list((await db.scalars(
        select(Measurement).where(Measurement.child_id == child_id).order_by(Measurement.measurement_date.desc())
    )).all())
    
    # Create Word document
    doc = Document()
    
    # Header
    title = doc.add_heading(f"Child Growth Monitoring Report", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Child Information
    doc.add_heading("Child Information", level=2)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.rows[0].cells[0].text = "Full Name"
    info_table.rows[0].cells[1].text = child.full_name
    info_table.rows[1].cells[0].text = "Date of Birth"
    info_table.rows[1].cells[1].text = child.birth_date.strftime("%B %d, %Y")
    info_table.rows[2].cells[0].text = "Age (months)"
    info_table.rows[2].cells[1].text = str(age_months(child.birth_date))
    info_table.rows[3].cells[0].text = "Sex"
    info_table.rows[3].cells[1].text = child.sex.value.title()
    info_table.rows[4].cells[0].text = "Guardian"
    info_table.rows[4].cells[1].text = child.guardian_name or "N/A"
    info_table.rows[5].cells[0].text = "Contact"
    info_table.rows[5].cells[1].text = child.contact_number or "N/A"
    
    # Latest Measurement
    if measurements:
        doc.add_heading("Latest Measurement", level=2)
        latest = measurements[0]
        
        meas_table = doc.add_table(rows=8, cols=2)
        meas_table.style = 'Light Grid Accent 1'
        
        meas_table.rows[0].cells[0].text = "Date"
        meas_table.rows[0].cells[1].text = latest.measurement_date.strftime("%B %d, %Y")
        meas_table.rows[1].cells[0].text = "Weight (kg)"
        meas_table.rows[1].cells[1].text = str(latest.weight_kg)
        meas_table.rows[2].cells[0].text = "Height (cm)"
        meas_table.rows[2].cells[1].text = str(latest.height_cm)
        meas_table.rows[3].cells[0].text = "MUAC (cm)"
        meas_table.rows[3].cells[1].text = str(latest.muac_cm or "N/A")
        meas_table.rows[4].cells[0].text = "Weight-for-Age (WAZ)"
        meas_table.rows[4].cells[1].text = str(latest.waz or "N/A")
        meas_table.rows[5].cells[0].text = "Height-for-Age (HAZ)"
        meas_table.rows[5].cells[1].text = str(latest.haz or "N/A")
        meas_table.rows[6].cells[0].text = "Weight-for-Height (WHZ)"
        meas_table.rows[6].cells[1].text = str(latest.whz or "N/A")
        meas_table.rows[7].cells[0].text = "Overall Status"
        meas_table.rows[7].cells[1].text = latest.overall_status.value.replace("_", " ").title() if latest.overall_status else "N/A"
        
        # Observations
        if latest.notes:
            doc.add_heading("Latest Observations", level=2)
            doc.add_paragraph(latest.notes)
    
    # Measurement History Table
    if measurements:
        doc.add_heading("Measurement History", level=2)
        hist_table = doc.add_table(rows=len(measurements) + 1, cols=8)
        hist_table.style = 'Light Grid Accent 1'
        
        header_cells = hist_table.rows[0].cells
        header_cells[0].text = "Date"
        header_cells[1].text = "Weight"
        header_cells[2].text = "Height"
        header_cells[3].text = "MUAC"
        header_cells[4].text = "WAZ"
        header_cells[5].text = "HAZ"
        header_cells[6].text = "WHZ"
        header_cells[7].text = "Status"
        
        for i, m in enumerate(measurements, 1):
            cells = hist_table.rows[i].cells
            cells[0].text = m.measurement_date.strftime("%Y-%m-%d")
            cells[1].text = str(m.weight_kg)
            cells[2].text = str(m.height_cm)
            cells[3].text = str(m.muac_cm or "")
            cells[4].text = str(m.waz or "")
            cells[5].text = str(m.haz or "")
            cells[6].text = str(m.whz or "")
            cells[7].text = m.overall_status.value.replace("_", " ").title() if m.overall_status else ""
    
    doc.add_paragraph()  # Spacing
    footer = doc.add_paragraph("Generated from HMS System")
    footer.style = 'Footer'
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save to BytesIO
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return StreamingResponse(
        iter([doc_bytes.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={child.full_name}_complete_report.docx"}
    )

